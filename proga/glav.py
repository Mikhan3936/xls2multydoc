#!/usr/bin/env python
# coding: utf-8

import numpy as np
import pandas as pa
import openpyxl as op
import datetime as DT
from docxtpl import DocxTemplate
from docx.shared import Cm
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from tkinter import *
#import spisok_test.py as spisok
def clicked():
    spisok(dat1_.get(),dat2_.get())
def spisok(da1,da2):

    def fios1(s): #функция возвращает укороченные фио
        try:fam=s.split()[0]
        except: fam="    "
        try:im=s.split()[1]
        except: im="    "
        try: ot=s.split()[2]
        except: ot="    "
        al433=fam[0:4]+im[0:3]+ot[0:3]
        return al433

    def damt(d):
        def default():
            print("Incorrect option")
        d1=d.split()[0]
        d2=d.split()[1]
        d3=d.split()[2]
        dict={
            "января":"1",
            "февраля":"2",
            "марта":"3",
            "апреля":"4",
            "мая":"5",
            "июня":"6",
            "июля":"7",
            "августа":"8",
            "сентября":"9",
            "октября":"10",
            "ноября":"11",
            "декабря":"12",
        }

        da=DT.datetime.strptime(d3+dict.get(d2,"1")+d1, '%Y%m%d').date()
        return da
    dat1="2021-02-01"
    dat2="2021-03-31"
    dat1=da1
    dat2=da2

    wb=op.open('sotrud.xlsx')
    ws = wb.active
    #Внутреннее совместительство - исключить из выборки
    col='D'  #вид работы
    col1='A' #ФИО
    col2='C' #подразделение
    row= 1
    maxrow=ws.max_row
    #sotrud=np.array()
    sotrud=np.empty((maxrow,3),dtype=object)
    while row < maxrow:
        if(ws['{0}{1}'.format(col, row)].value is not None):
             if (ws['{0}{1}'.format(col, row)].value!="Внутреннее совместительство"):
                sotrud[row,0]=ws['{0}{1}'.format(col1, row)].value
                sotrud[row,1]=ws['{0}{1}'.format(col2, row)].value
                sotrud[row,2]=fios1(sotrud[row,0])
        row=row+1
    #Заполнена база сотрудников с подразделениями sotrud(ФИО, подразделение)
    maxr=row
    #print(sotrud)
    wb.close()
    #string.find(substring,start,end) поиск в строке подстроки
    #Другой файл - база
    wb=op.open('basa.xlsx')
    ws = wb.active
    col='I'  #ФИО
    col7='A' #Москва
    col1='H' #Название программы
    col2='C' #Номер приказа
    col3='K' #Номер бланка
    col4='L' #Номер удостоверения
    col5='F' #число 
    col6='G' #месяц год
    row= 2
    maxrow=ws.max_row
    #sotrud=np.array()
    basa=np.empty((maxrow,8),dtype=object)
    while row < maxrow:
        if(ws['{0}{1}'.format(col, row)].value is not None):
            basa[row,0]=ws['{0}{1}'.format(col, row)].value #ФИО дательный
            basa[row,1]=ws['{0}{1}'.format(col1, row)].value #Название программы
            basa[row,2]=fios1(basa[row,0]) #ФИО433
            basa[row,3]=ws['{0}{1}'.format(col2, row)].value #Номер приказа
            basa[row,4]=ws['{0}{1}'.format(col3, row)].value #Номер бланка
            basa[row,5]=ws['{0}{1}'.format(col4, row)].value #Номер удостоверения
            basa[row,7]=ws['{0}{1}'.format(col7, row)].value #Москва
            basa[row,6]=damt(str(ws['{0}{1}'.format(col5, row)].value)+' '+str(ws['{0}{1}'.format(col6, row)].value)) # Дата
            #date = DT.datetime.strptime(text, '%Y%m%d').date()
        row=row+1
    #Заполнена база сотрудников с подразделениями sotrud(ФИО, подразделение)
    #print(basa)
    wb.close()
    #string.find(substring,start,end) поиск в строке подстроки
    #сортировка по датам
    row= 2
    seedr=-1 #Число записей выборки
    seed=np.empty((maxrow,7),dtype=object)
    while row < maxrow:
        if(basa[row,7]=='Москва' and basa[row,6]>=DT.datetime.date(DT.datetime.strptime(dat1, "%Y-%m-%d")))and(basa[row,6]<=DT.datetime.date(DT.datetime.strptime(dat2, "%Y-%m-%d"))):
            seedr=seedr+1
            seed[seedr,0]=basa[row,0]
            seed[seedr,1]=basa[row,1]
            seed[seedr,2]=basa[row,2]
            seed[seedr,3]=basa[row,3]
            seed[seedr,4]=basa[row,4]
            seed[seedr,5]=basa[row,5]
            seed[seedr,6]=basa[row,6]
        row=row+1    
    #print(seedr)

    # формирование списков подразделений и списка людей
    row=0
    sed=0
    #dep1=0
    depart=np.empty((seedr+1,7),dtype=object)# массив для значений по департаментам
    dep=np.empty(seedr+1,dtype='U255')# массив с названиями департаментов
    pro=np.empty(seedr+1,dtype='U255')# массив с названиями программ
    while sed < seedr:
        while row < maxr:
            if (seed[sed,2]==sotrud[row,2]):
                depart[sed,0]=seed[sed,0]#ФИО дательный
                depart[sed,1]=seed[sed,1]#Название программы
                pro[sed]=seed[sed,1]#Название программы
                depart[sed,2]=seed[sed,3]#Номер приказа
                depart[sed,3]=seed[sed,4]#Номер бланка
                depart[sed,4]=seed[sed,5]#Номер удостоверения
                depart[sed,5]=sotrud[row,0]# ФИО именительный
                #print(depart[sed,5])
                depart[sed,6]=sotrud[row,1]# Подразделение
                dep[sed]=sotrud[row,1]# Подразделение
            row=row+1
        sed=sed+1
        row=0
    #[x for x in dep if x != '']
    unq=np.unique(dep)  # департаменты уникальные      
    #print(unq)
    depac=len(unq) # число департаментов"
    unp=np.unique(pro) # программы уникальные
    proac=len(unp) # число программ в выборке
    #print(unq)
    #cell.add_paragraph(transfer['comment']) добавление в поле таблицы
    #row.cells[0].merge(row.cells[-1]) объединение ячеек
    dep_one=np.empty((seedr+1,7),dtype=object)# массив для значений по 1 департаменту
    pro_one=np.empty(depac,dtype='U255')# массив с названиями программ по 1 департаменту
    row=0
    row2=0
    dang=0
    onedep=0 #число сотрудников в 1 департаменте
    onde=0 #число программ в 1 департаменте
    #print(depac)
    #dtype1 = [('num', 'U4'), ('fio', 'U128'), ('numud', 'U50'),('numbl','U50'),('nump','U50'),('proga','U255')]
    #dep_t=np.empty((seedr+1,6),dtype=dtype1)#массив для вывода в файл
    while dang < depac: # перемещение по выборке
        #dep_one[:] #опустошение массива?
        dep_t=np.empty((seedr+1,6),dtype='U255')#массив для вывода в файл
        onedep_one=0
        row2=0
        row=0
        #запись файла, имя департамент с защитой от пустого значения
        #pro_one=np.empty(seedr+1,dtype='U255')
        if unq[dang]!='':
            #print(dang)
            while row < seedr: # перечисление по базе
                if depart[row,6]==unq[dang]: #сравнение с департаментом по выборке
                    dep_one[onedep,0]=depart[row,0]#ФИО дательный
                    dep_one[onedep,1]=depart[row,1]#Название программы
                    pro_one[onedep]=depart[row,1]#Название программы
                    dep_one[onedep,2]=depart[row,2]#Номер приказа
                    dep_one[onedep,3]=depart[row,3]#Номер бланка
                    dep_one[onedep,4]=depart[row,4]#Номер удостоверения
                    dep_one[onedep,5]=depart[row,5]# ФИО именительный
                    dep_one[onedep,6]=depart[row,6]# Подразделение
                    onedep=onedep+1
                row=row+1
            # сформированный список с людьми из 1 департамента
            # формирование списков программ по 1 департаменту

            onde1=np.unique(pro_one)#(pro_one)#уникальные названия программ в департаменте
            onde=len(onde1)
            ops=0
            #print(onedep)


            if(onde>1):
                filename=str('./1/'+unq[dang].replace('"',''))+'_'+dat1+'_'+dat2+'.docx' #имя файла str(dang)
                filename1=str('./1/1/'+unq[dang].replace('"',''))+'_'+dat1+'_'+dat2+'.docx' #имя файла str(dang)где1 слушатель
                #print(filename)
                doc = DocxTemplate("shablon.docx")# загрузка из шаблона
                onedep_one=onedep+onde #число строк таблицы с прогами и фио
                row3=0
                #п/п; фио; рег ном; удост ном; приказ
                table = doc.add_table(rows=1,cols=5, style='Table Grid')#(onedep_one-1,6)
                table.border = 1
                table.autofit = False 
                table.allow_autofit = False
                table.columns[0].width = Cm(1.0)
                table.rows[0].cells[0].width = Cm(1.0)
                table.columns[1].width = Cm(6.5)
                table.rows[0].cells[1].width = Cm(7.5)  
                table.columns[4].width = Cm(2.5)
                table.rows[0].cells[4].width = Cm(2.5) 
                #table.cols[0].width=Cm(1)
                while row3 < onde:# перечисление по программам 1 департамента
                    row2=0
                    #добавление нового человека
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '№ п/п'
                    hdr_cells[1].text = 'ФИО слушателя программы ДПО'
                    hdr_cells[2].text = 'Внутренний регистрационный номер'
                    hdr_cells[3].text = 'Серия и номер удостоверения'
                    hdr_cells[4].text = 'Номер приказа'                
                    hdr_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                    hdr_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    hdr_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
                    hdr_cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    hdr_cells[3].paragraphs[0].runs[0].font.size = Pt(12)
                    hdr_cells[3].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    hdr_cells[3].paragraphs[0].runs[0].font.size = Pt(12)
                    hdr_cells[4].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    hdr_cells[4].paragraphs[0].runs[0].font.size = Pt(12)
                    numb=1 #номер по порядку для программы
                    while row2 < onedep:#перечисление по сотрудникам департамента
                        if (onde1[row3]==dep_one[row2,1]):
                            dep_t[row2,0]=str(numb)# номер по порядку (переделать)
                            dep_t[row2,1]=dep_one[row2,5]#ФИО именительный
                            dep_t[row2,2]='ПК '+str(dep_one[row2,4])#Номер удостоверения
                            dep_t[row2,3]='06.03д3/'+str(dep_one[row2,3])#Номер бланка
                            dep_t[row2,4]=dep_one[row2,2]#Номер приказа
                            dep_t[row2,5]=onde1[row3]
                            #print(dep_t[row2:])
                            numb=numb+1
                            if(str(dep_t[row2,0])=="1"):
                                #print(dep_t[row2,0])
                                row_cells = table.add_row().cells
                                row_cells[0].text = str(dep_t[row2,5])
                                row_cells[0].merge(row_cells[-1])
                                row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                                row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)

                            row_cells = table.add_row().cells
                            row_cells[0].text = str(dep_t[row2,0])
                            row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                            row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                            row_cells[1].text = str(dep_t[row2,1])
                            row_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
                            row_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
                            row_cells[3].text = str(dep_t[row2,2])
                            row_cells[3].paragraphs[0].runs[0].font.name = 'Times New Roman'
                            row_cells[3].paragraphs[0].runs[0].font.size = Pt(12)
                            row_cells[2].text = str(dep_t[row2,3])
                            row_cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
                            row_cells[2].paragraphs[0].runs[0].font.size = Pt(12)
                            row_cells[4].text = str(dep_t[row2,4])
                            row_cells[4].paragraphs[0].runs[0].font.name = 'Times New Roman'
                            row_cells[4].paragraphs[0].runs[0].font.size = Pt(12)
                        row2=row2+1

                    row3=row3+1
                #np.sort(dep_t , order='proga')
                #dep_t.tolist()
                #print(dep_t)
                #dep_t.sort()
                #print(dep_t)    
                #добавление в файл filename таблицы с данными из dep_t
                style = doc.styles['norm']
                font = style.font
                font="Times New Roman"
                #doc.add_paragraph("",style='Normal')
                #table = doc.add_table(rows=1,cols=5)#(onedep_one-1,6)
                row_cells = table.add_row().cells

                row_cells[0].text = 'Всего выдано документов (шт.): '+str(onedep)
                row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                #row_cells[0].font.name='Times New Roman'
                row_cells[0].merge(row_cells[-1])
                doc.add_paragraph("",style='norm')
                doc.add_paragraph("«___»______________2021 г.                                    ___________ /______________________/", style="norm")
                context = { 'pro' : unq[dang],'tab':''}
                doc.render(context)
                fil=filename
                if(onedep==1):
                    fil=filename1 #если 1 слушатель то в отдельный каталог
                doc.save(fil)
                onedep=0

            onde1=[]
            proone=[]
        dep_t=[]    
                        #добавление таблицы
                        #формирование строк таблицы с людьми 1 программы 1 департамента
                        #формирование файла документа
        dang=dang+1 #следующий департамент


    #print(seed)
    # Запись файла со списком всех из выборки
    sed=-1
    wb=op.Workbook()
    ws = wb.active
    while sed < seedr-1:
        sed=sed+1
        ws.cell(sed+1,1).value=depart[sed,0]#ФИО дательный
        ws.cell(sed+1,2).value=depart[sed,1]#Название программы
        ws.cell(sed+1,3).value=depart[sed,2]#Номер приказа
        ws.cell(sed+1,4).value=depart[sed,3]#Номер бланка
        ws.cell(sed+1,5).value=depart[sed,4]#Номер удостоверения
        ws.cell(sed+1,6).value=depart[sed,5]# ФИО именительный
        ws.cell(sed+1,7).value=depart[sed,6]# Подразделение

    wb.save('test.xlsx')
    print('all')
window = Tk()
window.title ("Прога для формирования списков")
lbl=Label(window,text="Формирование списков \nиз базы удостоверений и \n базы слушателей")
lbl.grid(column=0,row=0)
window.geometry('400x400')
btn =Button(window,text="Формирование.", command=clicked)
btn.grid(column=0,row=1)
lbl1=Label(window,text="Дата начала поиска (2021-02-01)")

lbl1.grid(column=0,row=2)
dat1_=Entry(window,width=10)
dat1_.grid(column=1,row=2)
lbl2=Label(window,text="Дата конца поиска (2021-03-31)")
lbl2.grid(column=0,row=4)
dat2_=Entry(window,width=10)
dat2_.grid(column=1,row=4)
window.mainloop()


